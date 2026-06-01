"""Export multi-formato del libro: Wattpad (txt), Word (docx), EPUB3, PDF.

Tollerante al modello dati: ogni capitolo puo' avere `content` (stringa) e/o
`subchapters` (lista di {title, content}). Cosi' funziona sia con l'output
dell'orchestratore sia con strutture annidate.

EPUB e' generato in puro Python (zipfile) senza dipendenze fragili.
PDF usa reportlab se disponibile, altrimenti viene saltato con grazia.
"""
from __future__ import annotations

import html
import os
import re
import uuid
import zipfile
from datetime import datetime
from typing import Dict, Iterator, List, Optional, Tuple

from .logging_config import get_logger

log = get_logger(__name__)


def _safe_name(title: str) -> str:
    title = (title or "libro").strip()
    name = "".join(c for c in title if c.isalnum() or c in (" ", "-", "_")).strip()
    return (name or "libro").replace(" ", "_")[:80]


def iter_chapters(book: Dict) -> Iterator[Tuple[str, str]]:
    """Restituisce (titolo_capitolo, testo) gestendo content e/o subchapters."""
    for i, ch in enumerate(book.get("chapters", []), 1):
        title = ch.get("title") or f"Capitolo {i}"
        content = (ch.get("content") or "").strip()
        if not content and ch.get("subchapters"):
            parts: List[str] = []
            for sc in ch["subchapters"]:
                if sc.get("title"):
                    parts.append(f"## {sc['title']}")
                if sc.get("content"):
                    parts.append(sc["content"])
            content = "\n\n".join(parts).strip()
        yield title, content


def _paragraphs(text: str) -> List[str]:
    return [p.strip() for p in re.split(r"\n\s*\n", text or "") if p.strip()]


# --------------------------------------------------------------------------- #
# Wattpad (txt)
# --------------------------------------------------------------------------- #
def export_txt(book: Dict, out_dir: str = "exports/wattpad") -> str:
    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, f"{_safe_name(book.get('title'))}.txt")
    blocks = [book.get("title", ""), "", book.get("plot", ""), ""]
    for title, content in iter_chapters(book):
        blocks += [title, "", content, ""]
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(blocks))
    log.info("Export TXT/Wattpad: %s", path)
    return path


# --------------------------------------------------------------------------- #
# Word (docx) — opzionale (python-docx)
# --------------------------------------------------------------------------- #
def export_docx(book: Dict, out_dir: str = "exports/word", cover_path: Optional[str] = None) -> Optional[str]:
    try:
        import docx
        from docx.shared import Inches, Pt
    except ImportError:
        log.warning("python-docx non installato: salto export Word")
        return None

    doc = docx.Document()
    if cover_path and os.path.exists(cover_path):
        try:
            doc.add_picture(cover_path, width=Inches(6))
        except Exception as e:  # immagine non valida: non bloccare l'export
            log.warning("Copertina non aggiunta al docx: %s", e)
    doc.add_heading(book.get("title", "Senza Titolo"), 0)
    if book.get("plot"):
        p = doc.add_paragraph(book["plot"])
        p.runs[0].italic = True if p.runs else None
    for title, content in iter_chapters(book):
        doc.add_heading(title, level=1)
        for para in _paragraphs(content):
            doc.add_paragraph(para)

    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, f"{_safe_name(book.get('title'))}.docx")
    doc.save(path)
    log.info("Export Word: %s", path)
    return path


# --------------------------------------------------------------------------- #
# EPUB3 — puro Python (zipfile)
# --------------------------------------------------------------------------- #
def _xhtml(title: str, body_html: str) -> str:
    return (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="it">\n'
        f"<head><title>{html.escape(title)}</title>"
        '<meta charset="utf-8"/></head>\n'
        f"<body>\n{body_html}\n</body>\n</html>\n"
    )


def export_epub(book: Dict, out_dir: str = "exports/epub", cover_path: Optional[str] = None) -> str:
    os.makedirs(out_dir, exist_ok=True)
    title = book.get("title", "Senza Titolo")
    lang = book.get("language", "it")
    book_id = f"urn:uuid:{uuid.uuid4()}"
    path = os.path.join(out_dir, f"{_safe_name(title)}.epub")

    chapters = list(iter_chapters(book))
    chapter_files: List[Tuple[str, str]] = []  # (filename, title)

    cover_item = ""
    cover_manifest = ""
    cover_spine = ""
    cover_data = None
    cover_ext = None
    if cover_path and os.path.exists(cover_path):
        cover_ext = os.path.splitext(cover_path)[1].lstrip(".").lower() or "png"
        media = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png"}.get(cover_ext, "image/png")
        with open(cover_path, "rb") as f:
            cover_data = f.read()
        cover_manifest = (
            f'    <item id="cover-image" href="cover.{cover_ext}" media-type="{media}" properties="cover-image"/>\n'
            '    <item id="cover" href="cover.xhtml" media-type="application/xhtml+xml"/>\n'
        )
        cover_spine = '    <itemref idref="cover"/>\n'
        cover_item = _xhtml("Copertina", f'<div><img src="cover.{cover_ext}" alt="Copertina"/></div>')

    manifest_items, spine_items, nav_lis = [], [], []
    for i, (ch_title, content) in enumerate(chapters, 1):
        fname = f"chap_{i:03d}.xhtml"
        body = f"<h1>{html.escape(ch_title)}</h1>\n" + "\n".join(
            f"<p>{html.escape(p)}</p>" for p in _paragraphs(content)
        )
        chapter_files.append((fname, _xhtml(ch_title, body)))
        manifest_items.append(f'    <item id="chap{i}" href="{fname}" media-type="application/xhtml+xml"/>')
        spine_items.append(f'    <itemref idref="chap{i}"/>')
        nav_lis.append(f'        <li><a href="{fname}">{html.escape(ch_title)}</a></li>')

    nav_xhtml = _xhtml(
        "Indice",
        '<nav epub:type="toc" xmlns:epub="http://www.idpf.org/2007/ops" id="toc">\n'
        f"      <h1>Indice</h1>\n      <ol>\n{chr(10).join(nav_lis)}\n      </ol>\n    </nav>",
    )

    opf = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id">\n'
        '  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">\n'
        f"    <dc:identifier id=\"book-id\">{book_id}</dc:identifier>\n"
        f"    <dc:title>{html.escape(title)}</dc:title>\n"
        f"    <dc:language>{html.escape(lang)}</dc:language>\n"
        f"    <dc:creator>{html.escape(book.get('author', os.getenv('AUTHOR_NAME', '')))}</dc:creator>\n"
        f"    <meta property=\"dcterms:modified\">{datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')}</meta>\n"
        "  </metadata>\n"
        "  <manifest>\n"
        '    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>\n'
        f"{cover_manifest}"
        + "\n".join(manifest_items) + "\n"
        "  </manifest>\n"
        "  <spine>\n"
        f"{cover_spine}"
        + "\n".join(spine_items) + "\n"
        "  </spine>\n"
        "</package>\n"
    )

    container = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n'
        '  <rootfiles>\n'
        '    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>\n'
        '  </rootfiles>\n</container>\n'
    )

    with zipfile.ZipFile(path, "w") as z:
        # mimetype DEVE essere il primo file e non compresso
        z.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        z.writestr("META-INF/container.xml", container, compress_type=zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/content.opf", opf, compress_type=zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/nav.xhtml", nav_xhtml, compress_type=zipfile.ZIP_DEFLATED)
        if cover_data is not None:
            z.writestr(f"OEBPS/cover.{cover_ext}", cover_data, compress_type=zipfile.ZIP_DEFLATED)
            z.writestr("OEBPS/cover.xhtml", cover_item, compress_type=zipfile.ZIP_DEFLATED)
        for fname, xhtml in chapter_files:
            z.writestr(f"OEBPS/{fname}", xhtml, compress_type=zipfile.ZIP_DEFLATED)
    log.info("Export EPUB: %s", path)
    return path


# --------------------------------------------------------------------------- #
# PDF — opzionale (reportlab)
# --------------------------------------------------------------------------- #
def export_pdf(book: Dict, out_dir: str = "exports/pdf") -> Optional[str]:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        log.warning("reportlab non installato: salto export PDF (pip install reportlab)")
        return None

    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, f"{_safe_name(book.get('title'))}.pdf")
    styles = getSampleStyleSheet()
    story = [Paragraph(html.escape(book.get("title", "Senza Titolo")), styles["Title"]), Spacer(1, 18)]
    if book.get("plot"):
        story += [Paragraph(html.escape(book["plot"]), styles["Italic"]), Spacer(1, 12)]
    for title, content in iter_chapters(book):
        story += [PageBreak(), Paragraph(html.escape(title), styles["Heading1"]), Spacer(1, 8)]
        for para in _paragraphs(content):
            story += [Paragraph(html.escape(para), styles["BodyText"]), Spacer(1, 6)]
    SimpleDocTemplate(path, pagesize=A4).build(story)
    log.info("Export PDF: %s", path)
    return path


def export_all(book: Dict, out_dir: str = "exports", cover_path: Optional[str] = None) -> Dict[str, Optional[str]]:
    """Esegue tutti gli export disponibili e restituisce i percorsi prodotti."""
    return {
        "txt": export_txt(book, os.path.join(out_dir, "wattpad")),
        "docx": export_docx(book, os.path.join(out_dir, "word"), cover_path=cover_path),
        "epub": export_epub(book, os.path.join(out_dir, "epub"), cover_path=cover_path),
        "pdf": export_pdf(book, os.path.join(out_dir, "pdf")),
    }
